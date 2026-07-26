from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A3, landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch, mm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_ROOT = REPOSITORY_ROOT / "outputs" / "demo" / "BuildCrew_BC-2026-0142" / "approval"
OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUTPUT_ROOT / "SR-081_Substitution_Request.docx"
DRAWING_PDF_PATH = OUTPUT_ROOT / "P-401_Manufacturer_Drawings.pdf"
EVIDENCE_PDF_PATH = OUTPUT_ROOT / "BC-2026-0142_Evidence_Summary.pdf"

INK = "000000"
MUTED = "555555"
BORDER = "DADCE0"
GREEN = colors.HexColor("#248D50")
LIGHT_GREEN = colors.HexColor("#EAF8EF")
GRAY = colors.HexColor("#66736D")


def set_cell_width(cell, width_twips: int) -> None:
    cell.width = width_twips
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.append(cell_width)
    cell_width.set(qn("w:w"), str(width_twips))
    cell_width.set(qn("w:type"), "dxa")


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "bottom", "insideH"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)
        borders.append(element)
    for edge in ("start", "end", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def style_paragraph(paragraph, *, size: int = 11, color: str = INK, bold: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def configure_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, INK),
        ("Heading 2", 16, 18, 6, INK),
        ("Heading 3", 14, 16, 4, "434343"),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Product Substitution Request")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = document.add_paragraph("SR-081 · Mission Bay Data Center · P-401 Chilled Water Pump")
    style_paragraph(subtitle, size=11, color=MUTED)


def add_simple_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.add_run(value)
        style_paragraph(paragraph, size=9, color=MUTED, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.add_run(value)
            style_paragraph(paragraph, size=9)
    set_table_geometry(table, widths)
    set_table_borders(table)


def generate_docx() -> None:
    document = Document()
    configure_document_styles(document)
    add_title(document)

    opening = document.add_paragraph(
        "BuildCrew recommends replacing the delayed specified pump with the Armstrong 4030 4×3×10. "
        "The proposed unit satisfies all 31 hard requirements, arrives before the required-on-site date, "
        "and passes project BIM coordination with one minor 25 mm piping adjustment."
    )
    style_paragraph(opening)

    document.add_heading("Decision requested", level=1)
    decision = document.add_paragraph(
        "Approve this package for submission to the mechanical engineer. This internal approval authorizes "
        "engineering review and a temporary inventory hold only; it does not authorize a purchase order or "
        "a permanent project-model update."
    )
    style_paragraph(decision)

    add_simple_table(
        document,
        ["Technical", "Coordination", "Delivery", "Commercial"],
        [["31 / 31 requirements", "0 critical clashes", "Aug 19 · 9 days early", "$31,680 installed"]],
        [2340, 2340, 2340, 2340],
    )

    document.add_heading("Reason for substitution", level=1)
    paragraph = document.add_paragraph(
        "The specified pump is delayed by 77 days because of a manufacturer component shortage. "
        "The revised ship date falls after chilled-water commissioning and exposes the project critical path."
    )
    style_paragraph(paragraph)

    document.add_heading("Technical equivalence", level=1)
    add_simple_table(
        document,
        ["Requirement", "Specified", "Proposed", "Evidence", "Result"],
        [
            ["Design flow", "420 gpm", "420 gpm", "M-601 / P-401", "Pass"],
            ["Design head", "96 ft", "98 ft", "Data sheet / p. 4", "Pass"],
            ["Inlet connection", '4 in FLG', '4 in FLG', "Submittal / sheet 6", "Pass"],
            ["Outlet connection", '3 in FLG', '3 in FLG', "Submittal / sheet 6", "Pass"],
            ["Motor", "15 HP · 480 V", "15 HP · 480 V", "Data sheet / p. 5", "Pass"],
            ["Overall envelope", "≤ 1,900 mm", "1,780 mm", "Dimension sheet / 6", "Pass"],
            ["Certification", "UL listed", "UL listed", "UL certificate", "Pass"],
        ],
        [1700, 1600, 1600, 2860, 1600],
    )

    document.add_heading("BIM coordination findings", level=1)
    paragraph = document.add_paragraph(
        "BuildCrew generated an IFC4 and GLB object from verified manufacturer dimensions, placed it at "
        "P-401, and checked equipment envelope, connectors, maintenance clearance, and adjacent systems."
    )
    style_paragraph(paragraph)
    add_simple_table(
        document,
        ["Check", "Finding", "Required action"],
        [
            ["Hard clash", "None", "No structural change"],
            ["Motor clearance", "915 mm preserved", "No action"],
            ["Suction connection", "25 mm offset", "Provide 25 mm spool adjustment"],
            ["Installation access", "Pass", "No action"],
        ],
        [2200, 3300, 3860],
    )

    document.add_heading("Cost and schedule effect", level=1)
    paragraph = document.add_paragraph(
        "Total installed cost is $31,680, a $4,260 increase over the delayed specified unit. "
        "The recommended replacement is available for delivery on August 19, preserving the August 28 "
        "required-on-site milestone and avoiding a modeled 77-day procurement delay."
    )
    style_paragraph(paragraph)

    document.add_heading("Evidence and files", level=1)
    for line in (
        "Replacement BIM: candidate-c.ifc and candidate-c.glb",
        "Coordination issues: candidate-c-coordination.bcfzip",
        "Source map and confidence report: source-map.json and confidence-report.json",
        "Commercial comparison: BuildCrew_Quote_Leveling.xlsx",
        "Manufacturer drawing set: P-401_Manufacturer_Drawings.pdf",
    ):
        paragraph = document.add_paragraph(line, style="List Bullet")
        style_paragraph(paragraph)

    document.save(DOCX_PATH)


def draw_dimension(canvas_object: canvas.Canvas, start, end, label: str, offset: float = 10 * mm) -> None:
    x1, y1 = start
    x2, y2 = end
    canvas_object.setStrokeColor(GRAY)
    canvas_object.setFillColor(GRAY)
    canvas_object.setLineWidth(0.7)
    canvas_object.line(x1, y1 + offset, x2, y2 + offset)
    canvas_object.line(x1, y1 + offset - 3 * mm, x1, y1 + offset + 3 * mm)
    canvas_object.line(x2, y2 + offset - 3 * mm, x2, y2 + offset + 3 * mm)
    canvas_object.setFont("Helvetica", 8)
    canvas_object.drawCentredString((x1 + x2) / 2, y1 + offset + 2 * mm, label)


def title_block(canvas_object: canvas.Canvas, sheet: str, title: str) -> None:
    width, _ = landscape(A3)
    canvas_object.setStrokeColor(colors.HexColor("#CAD4CE"))
    canvas_object.rect(16 * mm, 12 * mm, width - 32 * mm, 22 * mm)
    canvas_object.setFont("Helvetica-Bold", 15)
    canvas_object.setFillColor(colors.HexColor("#17211C"))
    canvas_object.drawString(21 * mm, 24 * mm, "BuildCrew")
    canvas_object.setFont("Helvetica", 8)
    canvas_object.setFillColor(GRAY)
    canvas_object.drawString(21 * mm, 18 * mm, "MISSION BAY DATA CENTER · P-401 REPLACEMENT")
    canvas_object.setFont("Helvetica-Bold", 10)
    canvas_object.setFillColor(colors.HexColor("#17211C"))
    canvas_object.drawString(118 * mm, 24 * mm, title)
    canvas_object.setFont("Helvetica", 8)
    canvas_object.drawString(118 * mm, 18 * mm, "COORDINATION LOD · NOT FOR FABRICATION")
    canvas_object.drawRightString(width - 21 * mm, 24 * mm, f"SHEET {sheet}")
    canvas_object.drawRightString(width - 21 * mm, 18 * mm, "DEMO EVIDENCE · 2026-07-26")


def generate_manufacturer_drawing_pdf() -> None:
    width, height = landscape(A3)
    drawing = canvas.Canvas(str(DRAWING_PDF_PATH), pagesize=landscape(A3))
    drawing.setTitle("P-401 Manufacturer Drawings")

    for sheet, title, view in (
        ("M-601.1", "PUMP ASSEMBLY · SIDE VIEW", "side"),
        ("M-601.2", "PUMP ASSEMBLY · FRONT VIEW", "front"),
        ("M-601.3", "PUMP ASSEMBLY · PLAN VIEW", "plan"),
    ):
        drawing.setFillColor(colors.white)
        drawing.rect(0, 0, width, height, fill=1, stroke=0)
        drawing.setStrokeColor(colors.HexColor("#E2E8E4"))
        for grid_x in range(20, int(width / mm) - 20, 10):
            drawing.line(grid_x * mm, 40 * mm, grid_x * mm, height - 16 * mm)
        for grid_y in range(40, int(height / mm) - 10, 10):
            drawing.line(16 * mm, grid_y * mm, width - 16 * mm, grid_y * mm)

        center_x = width * 0.50
        center_y = height * 0.56
        drawing.setStrokeColor(colors.HexColor("#17211C"))
        drawing.setLineWidth(1.3)
        if view == "side":
            drawing.roundRect(center_x - 105 * mm, center_y - 21 * mm, 90 * mm, 42 * mm, 12 * mm)
            drawing.circle(center_x + 16 * mm, center_y, 27 * mm)
            drawing.rect(center_x - 112 * mm, center_y - 29 * mm, 170 * mm, 7 * mm)
            drawing.circle(center_x + 49 * mm, center_y, 10 * mm)
            drawing.circle(center_x + 16 * mm, center_y + 37 * mm, 9 * mm)
            draw_dimension(drawing, (center_x - 112 * mm, center_y - 29 * mm), (center_x + 58 * mm, center_y - 29 * mm), "1,840 mm", -18 * mm)
            draw_dimension(drawing, (center_x - 112 * mm, center_y + 40 * mm), (center_x + 58 * mm, center_y + 40 * mm), "1,780 mm", 7 * mm)
        elif view == "front":
            drawing.circle(center_x, center_y, 39 * mm)
            drawing.circle(center_x, center_y, 11 * mm)
            drawing.rect(center_x - 48 * mm, center_y - 52 * mm, 96 * mm, 9 * mm)
            drawing.circle(center_x, center_y + 52 * mm, 10 * mm)
            draw_dimension(drawing, (center_x - 48 * mm, center_y - 52 * mm), (center_x + 48 * mm, center_y - 52 * mm), "780 mm", -15 * mm)
        else:
            drawing.roundRect(center_x - 105 * mm, center_y - 30 * mm, 90 * mm, 60 * mm, 12 * mm)
            drawing.circle(center_x + 16 * mm, center_y, 29 * mm)
            drawing.rect(center_x - 112 * mm, center_y - 38 * mm, 170 * mm, 76 * mm)
            draw_dimension(drawing, (center_x - 112 * mm, center_y - 38 * mm), (center_x + 58 * mm, center_y - 38 * mm), "1,840 mm", -15 * mm)

        drawing.setFillColor(LIGHT_GREEN)
        drawing.setStrokeColor(GREEN)
        drawing.roundRect(20 * mm, height - 42 * mm, 72 * mm, 18 * mm, 3 * mm, fill=1, stroke=1)
        drawing.setFillColor(GREEN)
        drawing.setFont("Helvetica-Bold", 9)
        drawing.drawString(25 * mm, height - 32 * mm, "VERIFIED DIMENSIONS")
        drawing.setFont("Helvetica", 7)
        drawing.drawString(25 * mm, height - 38 * mm, "Source: Armstrong demo submittal")
        title_block(drawing, sheet, title)
        drawing.showPage()
    drawing.save()


def generate_evidence_summary_pdf() -> None:
    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(
        str(EVIDENCE_PDF_PATH),
        pagesize=letter,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title="BuildCrew Evidence Summary",
    )
    title_style = ParagraphStyle(
        "BuildCrewTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=27,
        textColor=colors.HexColor("#17211C"),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        spaceAfter=14,
    )
    body_style = ParagraphStyle(
        "BuildCrewBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#2E3D36"),
    )
    story = [
        Paragraph("BuildCrew evidence summary", title_style),
        Paragraph(
            "BC-2026-0142 · Mission Bay Data Center · P-401 Chilled Water Pump",
            body_style,
        ),
        Spacer(1, 14),
    ]
    data = [
        ["Claim", "Verified value", "Source", "Grade", "Confidence"],
        ["Design flow", "420 gpm", "M-601 / P-401", "A", "100%"],
        ["Pump envelope", "1,780 × 760 × 980 mm", "Submittal / sheet 6", "A", "99%"],
        ["Motor clearance", "915 mm", "Installation manual / p.18", "A", "98%"],
        ["Inventory", "2 units · Aug 19", "Supplier quote 7614", "B", "100%"],
        ["Installed cost", "$31,680", "Quote leveling workbook", "A", "100%"],
    ]
    table = Table(data, colWidths=[1.4 * inch, 1.45 * inch, 1.7 * inch, 0.55 * inch, 0.8 * inch])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#17211C")),
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT_GREEN),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 18))
    story.append(
        Paragraph(
            "<b>Decision rule.</b> Installation-critical geometry requires grade A/B evidence and confidence of at least 95%. "
            "The BuildCrew BIM Engine blocks export when this threshold is not met.",
            body_style,
        )
    )
    document.build(story)


if __name__ == "__main__":
    generate_docx()
    generate_manufacturer_drawing_pdf()
    generate_evidence_summary_pdf()
    print(DOCX_PATH)
    print(DRAWING_PDF_PATH)
    print(EVIDENCE_PDF_PATH)
